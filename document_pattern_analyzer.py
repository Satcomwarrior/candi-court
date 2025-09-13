#!/usr/bin/env python3
"""
Document Pattern Analysis Tool for Legal Cases
Analyzes recent case filings for structure and language patterns to emulate successful patterns.
Uses NLP techniques to identify effective legal language and document structures.
"""

import os
import json
import spacy
from textblob import TextBlob
from pathlib import Path
from docx import Document
import re

# Load spacy model
try:
    nlp = spacy.load("en_core_web_sm")
    print("[NLP] spaCy model loaded for document analysis")
except OSError:
    print("[NLP] Warning: spaCy model not found. Some analysis features will be limited.")
    nlp = None

class LegalDocumentAnalyzer:
    """
    Analyzes legal documents to identify successful patterns and structures.
    """
    
    def __init__(self):
        self.legal_terms = [
            "RCW", "pursuant to", "wherefore", "heretofore", "whereas", "prima facie",
            "res judicata", "inter alia", "pro se", "ex parte", "in camera", "sua sponte",
            "order to show cause", "motion", "petition", "declaration", "affidavit",
            "service of process", "due process", "burden of proof", "preponderance",
            "clear and convincing", "beyond reasonable doubt", "material fact",
            "genuine issue", "summary judgment", "temporary restraining order",
            "preliminary injunction", "protection order", "no contact order"
        ]
        
        self.effective_legal_phrases = [
            "the evidence clearly demonstrates",
            "the record establishes",
            "the undisputed facts show",
            "the court should find",
            "justice requires",
            "the law mandates",
            "public policy supports",
            "the interests of justice",
            "fundamental fairness",
            "due process requires"
        ]
        
        self.document_types = {
            "motion": ["motion", "moving", "relief sought", "comes now"],
            "petition": ["petition", "petitioner", "relief requested", "pray"],
            "declaration": ["declaration", "declare", "personal knowledge", "attached"],
            "order": ["order", "ordered", "adjudged", "decreed", "it is hereby"],
            "response": ["response", "respondent", "denies", "admits", "lacks knowledge"]
        }

    def analyze_document_file(self, file_path):
        """
        Analyze a single legal document file for patterns and structure.
        """
        try:
            if file_path.endswith('.docx'):
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                
                analysis = {
                    "file_path": file_path,
                    "document_type": self.identify_document_type(text),
                    "structure_analysis": self.analyze_structure(doc),
                    "language_patterns": self.analyze_language_patterns(text),
                    "legal_effectiveness": self.assess_legal_effectiveness(text),
                    "formatting_analysis": self.analyze_formatting(doc),
                    "compliance_check": self.check_court_compliance(text)
                }
                
                return analysis
                
        except Exception as e:
            print(f"[ERROR] Failed to analyze {file_path}: {e}")
            return None

    def identify_document_type(self, text):
        """
        Identify the type of legal document based on content patterns.
        """
        text_lower = text.lower()
        scores = {}
        
        for doc_type, keywords in self.document_types.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            scores[doc_type] = score
            
        # Return the document type with highest score
        if scores:
            return max(scores, key=scores.get)
        return "unknown"

    def analyze_structure(self, doc):
        """
        Analyze document structure including headings, paragraphs, and organization.
        """
        structure = {
            "total_paragraphs": len(doc.paragraphs),
            "headings": [],
            "paragraph_lengths": [],
            "structure_score": 0
        }
        
        for para in doc.paragraphs:
            # Analyze headings
            if para.style.name.startswith('Heading'):
                structure["headings"].append({
                    "level": para.style.name,
                    "text": para.text,
                    "position": len(structure["paragraph_lengths"])
                })
            
            # Track paragraph lengths
            structure["paragraph_lengths"].append(len(para.text.split()))
        
        # Calculate structure score based on organization
        structure["structure_score"] = self.calculate_structure_score(structure)
        
        return structure

    def calculate_structure_score(self, structure):
        """
        Calculate a score for document structure quality.
        """
        score = 0
        
        # Points for having headings
        if structure["headings"]:
            score += len(structure["headings"]) * 10
        
        # Points for consistent paragraph lengths
        avg_length = sum(structure["paragraph_lengths"]) / len(structure["paragraph_lengths"])
        if 50 <= avg_length <= 200:  # Ideal paragraph length
            score += 20
        
        # Points for balanced structure
        if 5 <= structure["total_paragraphs"] <= 20:  # Reasonable document length
            score += 15
            
        return min(score, 100)  # Cap at 100

    def analyze_language_patterns(self, text):
        """
        Analyze language patterns for legal effectiveness.
        """
        patterns = {
            "legal_terminology_density": 0,
            "effective_phrases_used": [],
            "sentence_complexity": {},
            "formality_score": 0,
            "clarity_indicators": {}
        }
        
        # Calculate legal terminology density
        word_count = len(text.split())
        legal_term_count = sum(1 for term in self.legal_terms if term.lower() in text.lower())
        patterns["legal_terminology_density"] = (legal_term_count / word_count * 100) if word_count > 0 else 0
        
        # Find effective phrases
        for phrase in self.effective_legal_phrases:
            if phrase.lower() in text.lower():
                patterns["effective_phrases_used"].append(phrase)
        
        # Analyze sentence complexity using spaCy if available
        if nlp:
            doc = nlp(text)
            sentences = list(doc.sents)
            
            if sentences:
                avg_length = sum(len(sent.text.split()) for sent in sentences) / len(sentences)
                patterns["sentence_complexity"] = {
                    "average_sentence_length": avg_length,
                    "total_sentences": len(sentences),
                    "complexity_score": "high" if avg_length > 25 else "medium" if avg_length > 15 else "low"
                }
        
        # Calculate formality score using TextBlob
        blob = TextBlob(text)
        subjectivity = blob.sentiment.subjectivity
        patterns["formality_score"] = (1 - subjectivity) * 100  # More objective = more formal
        
        return patterns

    def assess_legal_effectiveness(self, text):
        """
        Assess the legal effectiveness of the document language.
        """
        effectiveness = {
            "persuasive_elements": [],
            "authority_citations": 0,
            "factual_assertions": 0,
            "legal_conclusions": 0,
            "overall_score": 0
        }
        
        text_lower = text.lower()
        
        # Look for persuasive elements
        persuasive_markers = [
            "clearly", "undoubtedly", "unequivocally", "overwhelmingly",
            "compelling", "substantial", "significant", "critical"
        ]
        
        for marker in persuasive_markers:
            if marker in text_lower:
                effectiveness["persuasive_elements"].append(marker)
        
        # Count authority citations (basic pattern matching)
        effectiveness["authority_citations"] = len(re.findall(r'\bRCW\s+\d+', text))
        effectiveness["authority_citations"] += len(re.findall(r'\b\w+\s+v\.\s+\w+', text))
        
        # Count factual assertions and legal conclusions
        fact_indicators = ["evidence shows", "the record reflects", "documentation establishes"]
        legal_indicators = ["the law requires", "the court must", "therefore"]
        
        for indicator in fact_indicators:
            effectiveness["factual_assertions"] += text_lower.count(indicator)
        
        for indicator in legal_indicators:
            effectiveness["legal_conclusions"] += text_lower.count(indicator)
        
        # Calculate overall effectiveness score
        score = len(effectiveness["persuasive_elements"]) * 5
        score += effectiveness["authority_citations"] * 10
        score += effectiveness["factual_assertions"] * 3
        score += effectiveness["legal_conclusions"] * 5
        
        effectiveness["overall_score"] = min(score, 100)
        
        return effectiveness

    def analyze_formatting(self, doc):
        """
        Analyze document formatting for professional presentation.
        """
        formatting = {
            "consistent_styling": True,
            "proper_spacing": True,
            "professional_appearance": 0
        }
        
        # This is a simplified analysis - in practice, you'd check
        # font consistency, spacing, alignment, etc.
        formatting["professional_appearance"] = 75  # Default score
        
        return formatting

    def check_court_compliance(self, text):
        """
        Check for compliance with common court requirements.
        """
        compliance = {
            "required_elements": [],
            "missing_elements": [],
            "compliance_score": 0
        }
        
        # Common required elements in family law documents
        required_elements = [
            ("case_number", r"case\s+no\.?\s*\d+", "Case number"),
            ("court_name", r"superior\s+court", "Court name"),
            ("parties", r"petitioner|respondent", "Party identification"),
            ("relief_sought", r"relief|prayer|wherefore", "Relief requested"),
            ("signature_line", r"signature|signed|declares", "Signature block")
        ]
        
        text_lower = text.lower()
        
        for element_id, pattern, description in required_elements:
            if re.search(pattern, text_lower):
                compliance["required_elements"].append(description)
            else:
                compliance["missing_elements"].append(description)
        
        # Calculate compliance score
        total_elements = len(required_elements)
        found_elements = len(compliance["required_elements"])
        compliance["compliance_score"] = (found_elements / total_elements * 100) if total_elements > 0 else 0
        
        return compliance

    def analyze_directory(self, directory_path):
        """
        Analyze all legal documents in a directory.
        """
        results = {
            "total_documents": 0,
            "successful_analyses": 0,
            "document_analyses": [],
            "patterns_summary": {},
            "recommendations": []
        }
        
        directory = Path(directory_path)
        if not directory.exists():
            print(f"[ERROR] Directory not found: {directory_path}")
            return results
        
        # Find all document files
        doc_files = list(directory.glob("*.docx")) + list(directory.glob("*.doc"))
        results["total_documents"] = len(doc_files)
        
        print(f"[INFO] Analyzing {len(doc_files)} documents in {directory_path}")
        
        for doc_file in doc_files:
            analysis = self.analyze_document_file(str(doc_file))
            if analysis:
                results["document_analyses"].append(analysis)
                results["successful_analyses"] += 1
        
        # Generate patterns summary
        results["patterns_summary"] = self.generate_patterns_summary(results["document_analyses"])
        
        # Generate recommendations
        results["recommendations"] = self.generate_recommendations(results["patterns_summary"])
        
        return results

    def generate_patterns_summary(self, analyses):
        """
        Generate summary of patterns across all analyzed documents.
        """
        if not analyses:
            return {}
        
        summary = {
            "most_common_doc_type": "",
            "average_structure_score": 0,
            "most_effective_phrases": [],
            "common_compliance_issues": [],
            "best_practices_identified": []
        }
        
        # Find most common document type
        doc_types = [analysis["document_type"] for analysis in analyses]
        summary["most_common_doc_type"] = max(set(doc_types), key=doc_types.count)
        
        # Calculate average structure score
        structure_scores = [analysis["structure_analysis"]["structure_score"] for analysis in analyses]
        summary["average_structure_score"] = sum(structure_scores) / len(structure_scores)
        
        # Find most effective phrases
        all_phrases = []
        for analysis in analyses:
            all_phrases.extend(analysis["language_patterns"]["effective_phrases_used"])
        phrase_counts = {phrase: all_phrases.count(phrase) for phrase in set(all_phrases)}
        summary["most_effective_phrases"] = sorted(phrase_counts.items(), key=lambda x: x[1], reverse=True)[:5]
        
        # Identify common compliance issues
        all_missing = []
        for analysis in analyses:
            all_missing.extend(analysis["compliance_check"]["missing_elements"])
        missing_counts = {item: all_missing.count(item) for item in set(all_missing)}
        summary["common_compliance_issues"] = sorted(missing_counts.items(), key=lambda x: x[1], reverse=True)[:3]
        
        return summary

    def generate_recommendations(self, patterns_summary):
        """
        Generate recommendations based on pattern analysis.
        """
        recommendations = []
        
        if patterns_summary.get("average_structure_score", 0) < 70:
            recommendations.append("Improve document structure by adding clear headings and maintaining consistent paragraph lengths.")
        
        if patterns_summary.get("common_compliance_issues"):
            missing_elements = [item[0] for item in patterns_summary["common_compliance_issues"]]
            recommendations.append(f"Ensure all documents include: {', '.join(missing_elements)}")
        
        if patterns_summary.get("most_effective_phrases"):
            top_phrases = [phrase[0] for phrase in patterns_summary["most_effective_phrases"][:3]]
            recommendations.append(f"Consider using these effective phrases: {'; '.join(top_phrases)}")
        
        recommendations.append("Use official WA court forms when available for better compliance.")
        recommendations.append("Include specific legal citations (RCW references) to strengthen arguments.")
        
        return recommendations

def main():
    """
    Main function to run document pattern analysis.
    """
    analyzer = LegalDocumentAnalyzer()
    
    # Analyze existing case documents
    print("[ANALYSIS] Starting Legal Document Pattern Analysis...")
    
    # Check for existing documents to analyze
    directories_to_analyze = [
        ".",  # Current directory
        "outputs",  # Generated documents
        "templates",  # Template documents
    ]
    
    all_results = {}
    
    for directory in directories_to_analyze:
        if os.path.exists(directory):
            print(f"\n[ANALYSIS] Analyzing directory: {directory}")
            results = analyzer.analyze_directory(directory)
            all_results[directory] = results
            
            # Display results
            print(f"[RESULTS] Analyzed {results['successful_analyses']}/{results['total_documents']} documents")
            
            if results["patterns_summary"]:
                summary = results["patterns_summary"]
                print(f"[PATTERNS] Most common document type: {summary.get('most_common_doc_type', 'Unknown')}")
                print(f"[PATTERNS] Average structure score: {summary.get('average_structure_score', 0):.1f}/100")
                
                if summary.get("most_effective_phrases"):
                    print("[PATTERNS] Most effective phrases found:")
                    for phrase, count in summary["most_effective_phrases"]:
                        print(f"  - '{phrase}' (used {count} times)")
            
            if results["recommendations"]:
                print("[RECOMMENDATIONS]")
                for rec in results["recommendations"]:
                    print(f"  - {rec}")
    
    # Save comprehensive analysis report
    report_file = "document_pattern_analysis_report.json"
    with open(report_file, 'w') as f:
        json.dump(all_results, f, indent=2)
    
    print(f"\n[COMPLETE] Document pattern analysis complete. Report saved to: {report_file}")

if __name__ == "__main__":
    main()