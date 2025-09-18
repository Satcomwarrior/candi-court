#!/usr/bin/env python3
"""
Generate court-ready documents from Snohomish County templates.

Inputs
- JSON file mapping placeholders to values (see case_data.sample.json)
- Snohomish templates created by the forms downloader:
  templates/family_law_forms/snohomish_county/

Usage
  python scripts/generate_case_documents.py --input case_data.json
  python scripts/generate_case_documents.py --input case_data.json --only motion declaration
  python scripts/generate_case_documents.py --input case_data.json --output-dir outputs/my_docs

If templates are missing, run:  python run_all.py forms
"""

import argparse
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List

try:
    from docx import Document
except ImportError as e:
    raise SystemExit("python-docx is required. Install with: pip install python-docx")


TEMPLATE_DIR = Path("templates/family_law_forms/snohomish_county")
TEMPLATES = {
    "motion": TEMPLATE_DIR / "snohomish_motion_template.docx",
    "declaration": TEMPLATE_DIR / "snohomish_declaration_template.docx",
    "contempt": TEMPLATE_DIR / "snohomish_contempt_motion_template.docx",
}


def replace_in_paragraphs(doc: Document, mapping: Dict[str, str]) -> None:
    for para in doc.paragraphs:
        text = para.text
        replaced = False
        for key, value in mapping.items():
            ph = f"{{{{{key}}}}}"
            if ph in text:
                text = text.replace(ph, str(value))
                replaced = True
        if replaced:
            para.text = text


def replace_in_tables(doc: Document, mapping: Dict[str, str]) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    replaced = False
                    for key, value in mapping.items():
                        ph = f"{{{{{key}}}}}"
                        if ph in text:
                            text = text.replace(ph, str(value))
                            replaced = True
                    if replaced:
                        para.text = text


def fill_template(template_path: Path, mapping: Dict[str, str], output_path: Path) -> Path:
    doc = Document(str(template_path))
    replace_in_paragraphs(doc, mapping)
    replace_in_tables(doc, mapping)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def ensure_templates_exist() -> None:
    missing = [name for name, p in TEMPLATES.items() if not p.exists()]
    if missing:
        raise SystemExit(
            "Missing Snohomish templates: "
            + ", ".join(missing)
            + "\nRun: python run_all.py forms"
        )


def load_case_data(path: Path) -> Dict[str, str]:
    data = json.loads(path.read_text(encoding="utf-8"))
    # Normalize keys to str
    return {str(k): ("" if v is None else str(v)) for k, v in data.items()}


def main() -> None:
    ap = argparse.ArgumentParser(description="Generate Snohomish legal documents from templates")
    ap.add_argument("--input", required=True, help="Path to JSON with placeholder values")
    ap.add_argument("--output-dir", default=str(Path("outputs") / "generated_docs"))
    ap.add_argument(
        "--only",
        nargs="*",
        choices=["motion", "declaration", "contempt"],
        help="Limit to these documents",
    )
    args = ap.parse_args()

    ensure_templates_exist()
    data = load_case_data(Path(args.input))

    when = datetime.now().strftime("%Y%m%d_%H%M")
    case_no = data.get("case_number", "CASE")
    outdir = Path(args.output_dir)

    targets: List[str] = args.only or ["motion", "declaration", "contempt"]
    outputs: List[Path] = []

    for name in targets:
        tpl = TEMPLATES[name]
        if name == "motion":
            fname = f"Motion_{case_no}_{when}.docx"
        elif name == "declaration":
            fname = f"Declaration_{case_no}_{when}.docx"
        else:
            fname = f"Contempt_Motion_{case_no}_{when}.docx"
        out = outdir / fname
        outp = fill_template(tpl, data, out)
        print(f"Generated: {outp}")
        outputs.append(outp)

    print(f"Done. {len(outputs)} document(s) created in {outdir}")


if __name__ == "__main__":
    main()

