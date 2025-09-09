#!/usr/bin/env python3
"""
Create a basic Word template for the legal workflow system.
"""

from docx import Document
import os

def create_basic_motion_template():
    """Create a basic motion template with placeholders."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('SUPERIOR COURT OF WASHINGTON FOR SNOHOMISH COUNTY', 0)
    title.alignment = 1  # Center alignment
    
    # Add case header
    doc.add_paragraph()
    case_header = doc.add_paragraph()
    case_header.add_run('In re: {{petitioner_name}} and {{respondent_name}}').bold = True
    
    doc.add_paragraph('Case No. {{case_number}}')
    doc.add_paragraph()
    
    # Add motion title
    motion_title = doc.add_heading('MOTION FOR [RELIEF REQUESTED]', level=1)
    motion_title.alignment = 1
    
    doc.add_paragraph()
    
    # Add introduction
    doc.add_paragraph('TO THE HONORABLE COURT:')
    doc.add_paragraph()
    
    # Add main content sections with placeholders
    doc.add_heading('I. INTRODUCTION', level=2)
    doc.add_paragraph('{{introduction_text}}')
    doc.add_paragraph()
    
    doc.add_heading('II. FACTS', level=2)
    doc.add_paragraph('{{facts}}')
    doc.add_paragraph()
    
    doc.add_heading('III. LEGAL ARGUMENTS', level=2)
    doc.add_paragraph('{{legal_arguments}}')
    doc.add_paragraph()
    
    doc.add_heading('IV. CONCLUSION', level=2)
    doc.add_paragraph('{{conclusion}}')
    doc.add_paragraph()
    
    # Add signature block
    doc.add_paragraph('Respectfully submitted,')
    doc.add_paragraph()
    doc.add_paragraph('_________________________')
    doc.add_paragraph('{{your_name}}')
    doc.add_paragraph('Pro Se Petitioner')
    doc.add_paragraph('{{your_address}}')
    doc.add_paragraph('{{your_phone}}')
    doc.add_paragraph('{{your_email}}')
    
    # Ensure templates directory exists
    os.makedirs('templates', exist_ok=True)
    
    # Save the template
    template_path = 'templates/motion_template.docx'
    doc.save(template_path)
    print(f"Basic motion template created at: {template_path}")
    
    return template_path

if __name__ == "__main__":
    create_basic_motion_template()