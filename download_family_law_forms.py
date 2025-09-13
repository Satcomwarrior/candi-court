#!/usr/bin/env python3
"""
Enhanced Family Law Forms Downloader for Washington State Courts

Automates the download of official blank legal forms and templates from:
- Washington Courts Forms Library
- Washington Law Help Form Library
- Snohomish County Superior Court specific forms

Ensures compliance with Snohomish County Superior Court rules and formatting.
"""

import os
import json
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple

import requests


class WashingtonFormsDownloader:
    """Enhanced downloader for Washington State legal forms."""

    def __init__(self, base_dir: str = None):
        if base_dir is None:
            base_dir = os.path.join(os.getcwd(), "templates", "family_law_forms")

        self.base_dir = Path(base_dir)
        self.base_dir.mkdir(parents=True, exist_ok=True)

        # Washington Courts base URLs
        self.wa_courts_base = "https://www.courts.wa.gov/forms/docs/"
        self.wa_law_help_base = "https://www.washingtonlawhelp.org/"

        # Session for connection reuse
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        })

        # Download log
        self.download_log: List[Dict[str, str]] = []

    def get_family_law_forms_list(self) -> Dict[str, List[Tuple[str, str]]]:
        """Get comprehensive list of family law forms to download."""
        return {
            "uccjea_forms": [
                ("FL UCCJEA 801", "FL_UCCJEA_801.doc"),
                ("FL UCCJEA 802", "FL_UCCJEA_802.doc"),
                ("FL UCCJEA 803", "FL_UCCJEA_803.doc"),
                ("FL UCCJEA 804", "FL_UCCJEA_804.doc"),
            ],
            "protection_orders": [
                ("FL All Family 160", "FL_All_Family_160.doc"),
                ("FL All Family 161", "FL_All_Family_161.doc"),
                ("FL All Family 166", "FL_All_Family_166.doc"),
                ("FL All Family 167", "FL_All_Family_167.doc"),
            ],
            "service_forms": [
                ("FL All Family 101", "FL_All_Family_101_Proof_of_Personal_Service.doc"),
                ("FL All Family 102", "FL_All_Family_102_Return_of_Service.doc"),
            ],
            "motions_orders": [
                ("FL All Family 135", "FL_All_Family_135.doc"),
                ("FL All Family 140", "FL_All_Family_140.doc"),
            ],
        }

    def download_form(self, form_name: str, filename: str, category: str = "general") -> bool:
        """Download a single form with error handling and logging."""
        url = self.wa_courts_base + filename
        category_dir = self.base_dir / category
        category_dir.mkdir(exist_ok=True)

        local_filename = filename.replace(".doc", f"_{datetime.now().strftime('%Y%m%d')}.doc")
        filepath = category_dir / local_filename

        try:
            print(f"Downloading {form_name}...")
            response = self.session.get(url, timeout=30)

            if response.status_code == 200:
                with open(filepath, "wb") as f:
                    f.write(response.content)

                self.download_log.append(
                    {
                        "form_name": form_name,
                        "filename": filename,
                        "local_path": str(filepath),
                        "category": category,
                        "status": "success",
                        "timestamp": datetime.now().isoformat(),
                        "size_bytes": str(len(response.content)),
                    }
                )

                print(f"Downloaded: {form_name}")
                return True
            else:
                print(f"Failed: {form_name} - HTTP {response.status_code}")
                self.download_log.append(
                    {
                        "form_name": form_name,
                        "filename": filename,
                        "category": category,
                        "status": "failed",
                        "error": f"HTTP {response.status_code}",
                        "timestamp": datetime.now().isoformat(),
                    }
                )
                return False

        except Exception as e:
            print(f"Error: {form_name} - {e}")
            self.download_log.append(
                {
                    "form_name": form_name,
                    "filename": filename,
                    "category": category,
                    "status": "error",
                    "error": str(e),
                    "timestamp": datetime.now().isoformat(),
                }
            )
            return False

    def download_all_forms(self) -> Dict[str, int]:
        """Download all family law forms organized by category."""
        forms_list = self.get_family_law_forms_list()
        stats = {"total": 0, "successful": 0, "failed": 0}

        print("Starting download of Washington State Family Law Forms...")
        print(f"Download directory: {self.base_dir}")
        print("=" * 60)

        for category, forms in forms_list.items():
            print(f"\nDownloading {category.replace('_', ' ').title()} forms:")
            print("-" * 40)

            for form_name, filename in forms:
                stats["total"] += 1
                if self.download_form(form_name, filename, category):
                    stats["successful"] += 1
                else:
                    stats["failed"] += 1
                time.sleep(0.5)

        return stats

    def create_snohomish_county_templates(self):
        """Create Snohomish County specific templates with proper formatting."""
        print("\nCreating Snohomish County specific templates...")

        snohomish_dir = self.base_dir / "snohomish_county"
        snohomish_dir.mkdir(exist_ok=True)

        self._create_motion_template(snohomish_dir)
        self._create_declaration_template(snohomish_dir)
        self._create_contempt_motion_template(snohomish_dir)

        print("Snohomish County templates created")

    def _create_motion_template(self, output_dir: Path):
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header.add_run("SUPERIOR COURT OF WASHINGTON FOR SNOHOMISH COUNTY").bold = True

            doc.add_paragraph()

            case_table = doc.add_table(rows=3, cols=2)
            case_table.style = "Table Grid"
            case_table.cell(0, 0).text = "In re: {{petitioner_name}} and {{respondent_name}}"
            case_table.cell(1, 0).text = "Petitioner and Respondent"
            case_table.cell(2, 0).text = ""
            case_table.cell(0, 1).text = "Case No. {{case_number}}"
            case_table.cell(1, 1).text = "MOTION FOR {{relief_requested}}"
            case_table.cell(2, 1).text = "Note on Motion Docket: {{hearing_date}}"

            doc.add_paragraph()
            doc.add_paragraph("TO THE HONORABLE COURT:")
            doc.add_paragraph()

            doc.add_heading("I. INTRODUCTION", level=2)
            doc.add_paragraph("{{introduction_paragraph}}")

            doc.add_heading("II. STATEMENT OF FACTS", level=2)
            doc.add_paragraph("{{facts_section}}")

            doc.add_heading("III. LEGAL ARGUMENT", level=2)
            doc.add_paragraph("{{legal_arguments}}")

            doc.add_heading("IV. CONCLUSION", level=2)
            doc.add_paragraph("{{conclusion_paragraph}}")

            doc.add_paragraph()
            doc.add_paragraph("Respectfully submitted,")
            doc.add_paragraph()
            doc.add_paragraph("_________________________________")
            doc.add_paragraph("{{your_name}}")
            doc.add_paragraph("Pro Se {{party_designation}}")
            doc.add_paragraph("{{your_address}}")
            doc.add_paragraph("{{your_city_state_zip}}")
            doc.add_paragraph("Phone: {{your_phone}}")
            doc.add_paragraph("Email: {{your_email}}")

            template_path = output_dir / "snohomish_motion_template.docx"
            doc.save(template_path)
        except ImportError:
            print("Warning: python-docx not available for template creation")

    def _create_declaration_template(self, output_dir: Path):
        try:
            from docx import Document

            doc = Document()
            title = doc.add_heading("DECLARATION OF {{declarant_name}}", 0)
            title.alignment = 1
            doc.add_paragraph()
            doc.add_paragraph("I, {{declarant_name}}, declare as follows:")
            doc.add_paragraph()
            for i in range(1, 11):
                doc.add_paragraph(f"{i}. {{declaration_paragraph_{i}}}")
            doc.add_paragraph()
            doc.add_paragraph(
                "I declare under penalty of perjury under the laws of the State of Washington that the foregoing is true and correct."
            )
            doc.add_paragraph()
            doc.add_paragraph("DATED this _____ day of _____________, 2025.")
            doc.add_paragraph()
            doc.add_paragraph("_________________________________")
            doc.add_paragraph("{{declarant_name}}")
            template_path = output_dir / "snohomish_declaration_template.docx"
            doc.save(template_path)
        except ImportError:
            pass

    def _create_contempt_motion_template(self, output_dir: Path):
        try:
            from docx import Document

            doc = Document()
            title = doc.add_heading("MOTION FOR ORDER TO SHOW CAUSE FOR CONTEMPT", 0)
            title.alignment = 1
            doc.add_paragraph("TO THE HONORABLE COURT:")
            doc.add_paragraph()
            doc.add_paragraph(
                "Petitioner moves this Court for an Order directing Respondent to show cause why Respondent should not be held in contempt for violation of the Court's orders, and states:"
            )
            doc.add_paragraph()
            doc.add_heading("I. BACKGROUND", level=2)
            doc.add_paragraph("{{background_facts}}")
            doc.add_heading("II. VIOLATIONS", level=2)
            doc.add_paragraph("{{violation_details}}")
            doc.add_heading("III. RELIEF REQUESTED", level=2)
            doc.add_paragraph("{{relief_requested}}")
            template_path = output_dir / "snohomish_contempt_motion_template.docx"
            doc.save(template_path)
        except ImportError:
            pass

    def save_download_log(self):
        log_file = self.base_dir / f"download_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(log_file, "w", encoding="utf-8") as f:
            json.dump(
                {
                    "download_session": {
                        "timestamp": datetime.now().isoformat(),
                        "total_downloads": len(self.download_log),
                        "base_directory": str(self.base_dir),
                    },
                    "downloads": self.download_log,
                },
                f,
                indent=2,
            )
        print(f"\nDownload log saved to: {log_file}")

    def generate_summary_report(self, stats: Dict[str, int]):
        print("\n" + "=" * 60)
        print("DOWNLOAD SUMMARY REPORT")
        print("=" * 60)
        print(f"Total forms attempted: {stats['total']}")
        print(f"Successfully downloaded: {stats['successful']}")
        print(f"Failed downloads: {stats['failed']}")
        print(f"Success rate: {(stats['successful'] / stats['total'] * 100):.1f}%")
        print(f"Download directory: {self.base_dir}")
        if stats["failed"] > 0:
            print("\nFailed downloads:")
            for entry in self.download_log:
                if entry["status"] in ["failed", "error"]:
                    print(f"  - {entry['form_name']}: {entry.get('error', 'Unknown error')}")


def get_available_templates() -> List[Dict[str, str]]:
    """Return a list of available template files for use in document generation."""
    forms_dir = Path("templates/family_law_forms")
    if not forms_dir.exists():
        return []
    templates: List[Dict[str, str]] = []
    for template_file in forms_dir.rglob("*.doc*"):
        templates.append(
            {
                "name": template_file.stem,
                "path": str(template_file),
                "type": "official_wa_court_form",
            }
        )
    return templates


def main():
    print("Enhanced Washington State Family Law Forms Downloader")
    print("=" * 60)
    downloader = WashingtonFormsDownloader()
    stats = downloader.download_all_forms()
    downloader.create_snohomish_county_templates()
    downloader.save_download_log()
    downloader.generate_summary_report(stats)
    print("\nForms download and template creation completed!")
    print(f"Forms are organized in: {downloader.base_dir}")


if __name__ == "__main__":
    main()
