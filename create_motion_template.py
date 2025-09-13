from docx import Document

def create_motion_template(path):
    doc = Document()
    doc.add_heading('Legal Motion Header', 0)
    doc.save(path)

if __name__ == "__main__":
    create_motion_template('c:/Users/Muddm/Downloads/templates/motion_template.docx')
    doc.add_paragraph('Date: ______________________________')
    doc.save(path)

if __name__ == "__main__":
    create_motion_template('c:/Users/Muddm/Downloads/templates/motion_template.docx')
