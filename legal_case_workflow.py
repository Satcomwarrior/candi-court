"""
FULL AUTOMATED LEGAL CASE WORKFLOW SCRIPT

Author: William Miller
Date: 09/07/2025
Version: 1.0

Enhanced with NLP analysis and official template integration.
"""
import os
import json
from dataclasses import dataclass
import spacy
from textblob import TextBlob
from pathlib import Path

# Uncomment and install if needed:
# pip install python-docx openai requests spacy textblob

# Import for document generation
from docx import Document

# Import enhanced template downloader
try:
    from download_family_law_forms import get_available_templates
except ImportError:
    def get_available_templates():
        return []

# Load spacy model for advanced NLP analysis
try:
    nlp = spacy.load("en_core_web_sm")
    print("[NLP] spaCy model loaded successfully")
except OSError:
    print("[NLP] spaCy model not found. Install with: python -m spacy download en_core_web_sm")
    nlp = None

# Uncomment and configure your AI API
# import openai

# ==== CONFIGURATION ====

YOUR_NAME = "William Miller"
YOUR_PHONE = "206-226-2085"
YOUR_EMAIL = "wmiller@muddmonkiesinc.com"

# Your OpenAI API key here (replace with your key)
# openai.api_key = "your_openai_api_key"

# === ENHANCED NLP ANALYSIS FUNCTIONS ===


@dataclass
class CommunicationAnalysisResult:
    """Container for AI and NLP communication analysis results."""

    ai_analysis: str
    nlp_analysis: dict
    combined_summary: str

    def __post_init__(self):
        # Normalize default values to keep downstream code simple
        self.ai_analysis = self.ai_analysis or ""
        self.nlp_analysis = self.nlp_analysis or {}
        self.combined_summary = self.combined_summary or ""

    # Provide string-like behaviour so existing code that expects a string
    # (for preview printing, slicing, etc.) keeps working.
    def __str__(self):
        return self.combined_summary

    def __repr__(self):
        return f"CommunicationAnalysisResult(summary={self.combined_summary!r})"

    def __len__(self):
        return len(self.combined_summary)

    def __getitem__(self, item):
        return self.combined_summary[item]

    def to_dict(self):
        """Return a dictionary representation for compatibility."""
        return {
            "ai_analysis": self.ai_analysis,
            "nlp_analysis": self.nlp_analysis,
            "combined_summary": self.combined_summary,
        }

def analyze_communication_patterns_nlp(text):
    """
    Enhanced communication analysis using spaCy and TextBlob for detecting
    coercive control, manipulation, and psychological patterns.
    """
    analysis_result = {
        "sentiment": {},
        "coercive_patterns": [],
        "psychological_indicators": [],
        "linguistic_patterns": {},
        "entities": [],
        "summary": ""
    }
    
    # TextBlob sentiment analysis
    blob = TextBlob(text)
    analysis_result["sentiment"] = {
        "polarity": blob.sentiment.polarity,  # -1 (negative) to 1 (positive)
        "subjectivity": blob.sentiment.subjectivity,  # 0 (objective) to 1 (subjective)
        "interpretation": get_sentiment_interpretation(blob.sentiment.polarity, blob.sentiment.subjectivity)
    }
    
    # spaCy analysis if available
    if nlp:
        doc = nlp(text)
        
        # Extract named entities
        analysis_result["entities"] = [
            {"text": ent.text, "label": ent.label_, "description": spacy.explain(ent.label_)}
            for ent in doc.ents
        ]
        
        # Analyze linguistic patterns
        analysis_result["linguistic_patterns"] = analyze_linguistic_patterns(doc)
        
        # Detect coercive control patterns
        analysis_result["coercive_patterns"] = detect_coercive_patterns(doc, text)
        
        # Identify psychological indicators
        analysis_result["psychological_indicators"] = detect_psychological_indicators(doc, text)
    
    # Generate summary
    analysis_result["summary"] = generate_nlp_summary(analysis_result)
    
    return analysis_result

def get_sentiment_interpretation(polarity, subjectivity):
    """Interpret TextBlob sentiment scores for legal context."""
    sentiment_type = "neutral"
    if polarity > 0.1:
        sentiment_type = "positive"
    elif polarity < -0.1:
        sentiment_type = "negative"
    
    objectivity = "objective" if subjectivity < 0.5 else "subjective"
    
    return f"{sentiment_type} sentiment, {objectivity} tone"

def analyze_linguistic_patterns(doc):
    """Analyze linguistic patterns that may indicate coercive behavior."""
    patterns = {
        "imperative_sentences": 0,
        "question_frequency": 0,
        "personal_pronouns": {"first": 0, "second": 0, "third": 0},
        "modal_verbs": 0,
        "negative_words": 0
    }
    
    # Coercive control linguistic markers
    coercive_markers = [
        "must", "need to", "have to", "should", "supposed to", "better",
        "always", "never", "nothing", "everything", "fault", "blame",
        "crazy", "overreacting", "dramatic", "sensitive"
    ]
    
    for token in doc:
        # Count questions
        if token.text == "?":
            patterns["question_frequency"] += 1
            
        # Count personal pronouns
        if token.tag_ in ["PRP", "PRP$"]:
            if token.lower_ in ["i", "me", "my", "mine"]:
                patterns["personal_pronouns"]["first"] += 1
            elif token.lower_ in ["you", "your", "yours"]:
                patterns["personal_pronouns"]["second"] += 1
            else:
                patterns["personal_pronouns"]["third"] += 1
                
        # Count modal verbs (control indicators)
        if token.tag_ == "MD":
            patterns["modal_verbs"] += 1
            
        # Count negative sentiment words
        if token.sentiment < 0:
            patterns["negative_words"] += 1
    
    return patterns

def detect_coercive_patterns(doc, text):
    """Detect patterns associated with coercive control."""
    patterns = []
    text_lower = text.lower()
    
    # Control patterns
    control_phrases = [
        "you can't", "you won't", "you shouldn't", "you're not allowed",
        "i forbid", "you're not going", "you better not", "don't you dare",
        "you need permission", "ask me first", "you have no right"
    ]
    
    # Isolation patterns
    isolation_phrases = [
        "nobody likes you", "your friends don't care", "family doesn't want",
        "you have no one", "i'm all you have", "they're against you",
        "turn against", "choose between"
    ]
    
    # Intimidation patterns
    intimidation_phrases = [
        "you'll regret", "you'll pay", "consequences", "what happens next",
        "don't make me", "you know what", "remember what happened"
    ]
    
    # Gaslighting patterns
    gaslighting_phrases = [
        "you're crazy", "you're imagining", "that never happened",
        "you're overreacting", "you're too sensitive", "you're dramatic",
        "you're paranoid", "you're losing it"
    ]
    
    pattern_categories = [
        ("control", control_phrases),
        ("isolation", isolation_phrases),
        ("intimidation", intimidation_phrases),
        ("gaslighting", gaslighting_phrases)
    ]
    
    for category, phrases in pattern_categories:
        found_phrases = [phrase for phrase in phrases if phrase in text_lower]
        if found_phrases:
            patterns.append({
                "category": category,
                "phrases": found_phrases,
                "severity": len(found_phrases)
            })
    
    return patterns

def detect_psychological_indicators(doc, text):
    """Detect psychological indicators in communication."""
    indicators = []
    text_lower = text.lower()
    
    # Emotional manipulation indicators
    manipulation_terms = [
        "disappointed", "hurt", "betrayed", "after everything",
        "ungrateful", "selfish", "don't care", "don't love"
    ]
    
    # Victim-blaming language
    victim_blaming = [
        "your fault", "you made me", "you caused", "because of you",
        "you're responsible", "you brought this on", "you asked for"
    ]
    
    # Minimization language
    minimization = [
        "just kidding", "not that bad", "exaggerating", "making it up",
        "it's nothing", "get over it", "move on", "forget about it"
    ]
    
    indicator_categories = [
        ("emotional_manipulation", manipulation_terms),
        ("victim_blaming", victim_blaming),
        ("minimization", minimization)
    ]
    
    for category, terms in indicator_categories:
        found_terms = [term for term in terms if term in text_lower]
        if found_terms:
            indicators.append({
                "type": category,
                "terms": found_terms,
                "count": len(found_terms)
            })
    
    return indicators

def generate_nlp_summary(analysis_result):
    """Generate a comprehensive summary of NLP analysis for legal use."""
    summary_parts = []
    
    # Sentiment summary
    sentiment = analysis_result["sentiment"]
    summary_parts.append(f"Communication exhibits {sentiment['interpretation']} "
                         f"(polarity: {sentiment['polarity']:.2f}, subjectivity: {sentiment['subjectivity']:.2f})")
    
    # Coercive patterns summary
    if analysis_result["coercive_patterns"]:
        pattern_types = [p["category"] for p in analysis_result["coercive_patterns"]]
        summary_parts.append(f"Coercive control patterns detected: {', '.join(set(pattern_types))}")
    
    # Psychological indicators summary
    if analysis_result["psychological_indicators"]:
        indicator_types = [i["type"] for i in analysis_result["psychological_indicators"]]
        summary_parts.append(f"Psychological manipulation indicators: {', '.join(set(indicator_types))}")
    
    # Entity summary
    if analysis_result["entities"]:
        entity_types = list(set([e["label"] for e in analysis_result["entities"]]))
        summary_parts.append(f"Key entities identified: {', '.join(entity_types)}")
    
    return ". ".join(summary_parts) if summary_parts else "No significant patterns detected."

def analyze_document_structure(file_path):
    """
    Analyze the structure of legal documents to identify patterns and compliance issues.
    """
    try:
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            
            structure_analysis = {
                "paragraph_count": len(doc.paragraphs),
                "heading_structure": [],
                "formatting_patterns": {},
                "legal_citations": [],
                "standard_clauses": []
            }
            
            # Analyze headings and structure
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    structure_analysis["heading_structure"].append({
                        "level": para.style.name,
                        "text": para.text
                    })
                
                # Look for legal citations (basic pattern matching)
                text = para.text
                if any(pattern in text for pattern in ["RCW", "v.", "Case No.", "§"]):
                    structure_analysis["legal_citations"].append(text[:100])
            
            return structure_analysis
            
    except Exception as e:
        print(f"[ERROR] Document structure analysis failed: {e}")
        return None

# === STEP 1: ENHANCED DOCUMENT GENERATION ===

def generate_legal_document(template_path, case_data, output_path):
    """
    Generate a legal document from a Word template by replacing placeholders.
    Enhanced to work with official WA court forms.
    """
    try:
        doc = Document(template_path)

        # Enhanced replacement strategy for official forms
        for para in doc.paragraphs:
            for key, value in case_data.items():
                placeholder = f"{{{{{key}}}}}"  # e.g., {{case_number}}
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, str(value))

        # Also check tables (common in official forms)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in case_data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

        doc.save(output_path)
        print(f"[INFO] Generated legal document saved at {output_path}")
        return output_path
        
    except Exception as e:
        print(f"[ERROR] Document generation failed: {e}")
        return None

# === STEP 2: ENHANCED COMMUNICATIONS ANALYSIS ===

def analyze_communications_via_ai(text):
    """
    Enhanced communication analysis using both AI and NLP techniques.
    Combines Google AI Studio analysis with local spaCy/TextBlob processing.
    Returns a CommunicationAnalysisResult which behaves like a string summary
    but also exposes detailed NLP findings.
    """
    print("[AI] Starting enhanced communication analysis...")
    
    # First, perform local NLP analysis
    nlp_analysis = analyze_communication_patterns_nlp(text)
    
    # Prepare enhanced prompt incorporating NLP findings
    nlp_summary = nlp_analysis.get("summary", "No significant patterns detected.")
    coercive_patterns = nlp_analysis.get("coercive_patterns", [])
    sentiment_info = nlp_analysis.get("sentiment", {})
    
    prompt = f"""Analyze the following communications for coercive control, manipulation, and perceived burdensomeness patterns.

Local NLP Analysis Results:
- Sentiment: {sentiment_info.get('interpretation', 'Unknown')}
- Coercive patterns detected: {len(coercive_patterns)} categories
- Summary: {nlp_summary}

Communication Text:
{text}

Provide a detailed legal analysis focusing on:
1. Evidence of coercive control under RCW 7.105.010
2. Manipulation tactics and psychological abuse indicators
3. Patterns of isolation, intimidation, or financial control
4. Relevance to family law and domestic violence cases
5. Recommendations for legal action or protective measures

Consider the technical NLP analysis alongside contextual legal interpretation."""

    print("[AI] Sending enhanced analysis prompt to Google AI Studio...")
    
    try:
        # Use the call_studio_ai function for real AI analysis
        from ai_studio_code import call_studio_ai
        ai_response = call_studio_ai(prompt)

        return CommunicationAnalysisResult(
            ai_analysis=ai_response,
            nlp_analysis=nlp_analysis,
            combined_summary=f"AI Analysis: {ai_response}\n\nNLP Technical Analysis: {nlp_summary}"
        )

    except Exception as e:
        print(f"[AI] Error calling Google AI Studio: {e}")
        # Enhanced fallback using NLP analysis
        fallback_analysis = generate_enhanced_fallback_analysis(nlp_analysis, text)
        return CommunicationAnalysisResult(
            ai_analysis=fallback_analysis,
            nlp_analysis=nlp_analysis,
            combined_summary=f"Enhanced Analysis: {fallback_analysis}\n\nNLP Technical Analysis: {nlp_summary}"
        )

def generate_enhanced_fallback_analysis(nlp_analysis, text):
    """
    Generate enhanced fallback analysis using NLP results when AI is unavailable.
    """
    analysis_parts = []
    
    # Sentiment-based analysis
    sentiment = nlp_analysis.get("sentiment", {})
    polarity = sentiment.get("polarity", 0)
    
    if polarity < -0.3:
        analysis_parts.append("Communication exhibits strong negative sentiment consistent with hostile or abusive patterns.")
    elif polarity < -0.1:
        analysis_parts.append("Communication shows negative sentiment that may indicate conflict or coercive behavior.")
    
    # Coercive pattern analysis
    coercive_patterns = nlp_analysis.get("coercive_patterns", [])
    if coercive_patterns:
        pattern_descriptions = []
        for pattern in coercive_patterns:
            category = pattern["category"]
            severity = pattern["severity"]
            if category == "control":
                pattern_descriptions.append(f"Controlling language detected ({severity} instances)")
            elif category == "isolation":
                pattern_descriptions.append(f"Isolation tactics identified ({severity} instances)")
            elif category == "intimidation":
                pattern_descriptions.append(f"Intimidation patterns found ({severity} instances)")
            elif category == "gaslighting":
                pattern_descriptions.append(f"Gaslighting language detected ({severity} instances)")
        
        analysis_parts.append("Coercive control patterns evidenced by: " + "; ".join(pattern_descriptions))
    
    # Psychological indicators
    psych_indicators = nlp_analysis.get("psychological_indicators", [])
    if psych_indicators:
        indicator_types = [i["type"].replace("_", " ") for i in psych_indicators]
        analysis_parts.append(f"Psychological manipulation indicators: {', '.join(set(indicator_types))}")
    
    # Legal relevance
    if coercive_patterns or psych_indicators:
        analysis_parts.append("These patterns are relevant to RCW 7.105.010 (coercive control) and support protective order proceedings.")
    
    # Recommendations
    if len(coercive_patterns) >= 2:
        analysis_parts.append("RECOMMENDATION: Strong evidence for protection order based on multiple coercive control patterns.")
    elif coercive_patterns or psych_indicators:
        analysis_parts.append("RECOMMENDATION: Document these patterns as supporting evidence for family law proceedings.")
    
    return " ".join(analysis_parts) if analysis_parts else "Analysis indicates standard communication without significant coercive patterns."

# === STEP 3: FACTS SUMMARIZATION AND ARGUMENT GENERATION ===

def summarize_case_facts(case_facts):
    prompt = f"Summarize these legal facts clearly for a Snohomish County family law motion:\n{case_facts}"

    print("[AI] Summarizing case facts with Google AI Studio...")
    
    try:
        # Use the call_studio_ai function for real AI analysis
        from ai_studio_code import call_studio_ai
        summary = call_studio_ai(prompt)
        return summary
    except Exception as e:
        print(f"[AI] Error calling Google AI Studio: {e}")
        # Fallback to simulated response
        summary = ("[Summary] The defendant engaged in repeated violations of "
                   "court orders, including financial abuse and property interference.")
        return summary

def generate_legal_arguments(summary_text):
    prompt = f"Generate 3-5 bullet points for legal arguments supporting a contempt or protective order motion based on this summary:\n{summary_text}"

    print("[AI] Generating legal arguments with Google AI Studio...")
    
    try:
        # Use the call_studio_ai function for real AI analysis
        from ai_studio_code import call_studio_ai
        arguments_text = call_studio_ai(prompt)
        # Try to parse bullet points from response
        arguments = [line.strip('- ').strip() for line in arguments_text.split('\n') if line.strip().startswith('-') or line.strip().startswith('•')]
        if not arguments:
            # If no bullet points found, split by newlines and filter
            arguments = [line.strip() for line in arguments_text.split('\n') if line.strip() and len(line.strip()) > 10]
        return arguments[:5]  # Limit to 5 arguments
    except Exception as e:
        print(f"[AI] Error calling Google AI Studio: {e}")
        # Fallback to simulated response
        arguments = [
            "Defendant willfully violated the protection order repeatedly.",
            "Financial transactions show ongoing support by plaintiff despite abuse.",
            "Pattern of coercive control demonstrated by property exclusion.",
            "Insurance cancellation during medical crisis constitutes bad faith.",
            "Digital trespass evidences unauthorized surveillance."
        ]
        return arguments

# === STEP 4: AI STUDIO DATA EXTRACTION (Google AI Studio Integration) ===

def ai_studio_extract_metadata(doc_paths):
    """
    Extract document metadata and entity extraction using Google AI Studio.
    Integrates with the call_studio_ai function for real AI analysis.
    """
    print("[AI Studio] Extracting metadata from documents...")
    
    try:
        # Import the call_studio_ai function
        from ai_studio_code import call_studio_ai
        
        # Prepare the prompt for metadata extraction
        prompt = """
        Analyze the provided documents and extract the following metadata:
        1. Important dates (format: YYYY-MM-DD)
        2. Person names mentioned
        3. Document categories (Evidence, Financial Records, Communications, etc.)
        4. Financial summaries (amounts, transactions, etc.)
        
        Please return the information in JSON format with these keys:
        - dates: array of dates found
        - persons: array of person names
        - categories: array of document categories
        - financial_summaries: object with financial data
        
        Focus on legal case information, particularly for family law and financial analysis.
        """
        
        # Call Google AI Studio with the documents
        response = call_studio_ai(prompt, files=doc_paths)
        
        # Try to parse JSON response, fallback to simulated data if needed
        try:
            import json
            # Look for JSON in the response
            import re
            json_match = re.search(r'\{[^}]*\}', response, re.DOTALL)
            if json_match:
                extracted = json.loads(json_match.group())
                return extracted
        except:
            pass
            
        # Fallback to simulated data if AI response cannot be parsed
        print("[AI Studio] Using fallback data due to parsing issues")
        extracted = {
            "dates": ["2024-09-13", "2025-03-08", "2025-06-28"],
            "persons": ["William Miller", "Candi Brightwell"],
            "categories": ["Evidence", "Financial Records", "Communications"],
            "financial_summaries": {"PayPalTransfers2025": 7355.0}
        }
        return extracted
        
    except Exception as e:
        print(f"[AI Studio] Error: {e}")
        print("[AI Studio] Using simulated data due to error")
        # Fallback to simulated data
        extracted = {
            "dates": ["2024-09-13", "2025-03-08", "2025-06-28"],
            "persons": ["William Miller", "Candi Brightwell"],
            "categories": ["Evidence", "Financial Records", "Communications"],
            "financial_summaries": {"PayPalTransfers2025": 7355.0}
        }
        return extracted

# === STEP 5: EXPERT WITNESS OUTREACH EMAIL GENERATION ===

def generate_intake_email(expert_type, provider_name):
    templates = {
        "psychologist": f"""Subject: Expert Psychological Evaluation and Testimony Inquiry for Snohomish County Family Law Case

Dear {provider_name},

I am preparing a family law case involving committed intimate relationship dissolution, psychological coercion, and trauma. I seek an expert psychologist for evaluation, expert testimony, and reporting.

Please provide your experience in forensic evaluations, availability for remote or in-person consultations, fee structure, and engagement procedures.

Thank you for your time.

Best regards,
{YOUR_NAME}
Phone: {YOUR_PHONE}
Email: {YOUR_EMAIL}
""",
        "financial_forensic": f"""Subject: Request for Financial Forensics Consultation and Expert Witness Services

Dear {provider_name},

I am involved in a family law proceeding needing a financial forensic expert to assess business valuation, asset tracing, and related financial abuse issues.

Kindly share your experience with family law cases, availability, fees, and how to engage your services.

Thank you.

Sincerely,
{YOUR_NAME}
Phone: {YOUR_PHONE}
Email: {YOUR_EMAIL}
""",
        "domestic_violence": f"""Subject: Inquiry Regarding Domestic Violence Expert Services for Family Law Case

Dear {provider_name},

I require an expert witness in domestic violence and coercive control to assist with expert reports and testimony in a family law matter.

Please share your experience, availability for remote/in-person services, fee schedule, and intake process.

Thank you for your assistance.

Kind regards,
{YOUR_NAME}
Phone: {YOUR_PHONE}
Email: {YOUR_EMAIL}
"""
    }
    return templates.get(expert_type, "")

def send_email(email_text):
    """Simulate sending or saving email drafts."""
    print("\n=== EMAIL DRAFT ===\n")
    print(email_text)
    print("\n=== END EMAIL ===\n")

# === MAIN EXECUTION FUNCTION ===

def main():
    print("[SYSTEM] Starting Enhanced Legal Case Workflow...")
    
    # Step 0: Download official templates if not already available
    available_templates = get_available_templates()
    if not available_templates:
        print("[TEMPLATES] No official templates found. Downloading...")
        try:
            from download_family_law_forms import download_family_law_forms
            download_family_law_forms()
            available_templates = get_available_templates()
            print(f"[TEMPLATES] Downloaded {len(available_templates)} official templates")
        except Exception as e:
            print(f"[TEMPLATES] Error downloading templates: {e}")
    else:
        print(f"[TEMPLATES] Found {len(available_templates)} available templates")

    # Case facts & communications with enhanced data
    case_data = {
        "case_number": "2025-PA-000123",
        "petitioner_name": "William Miller",
        "respondent_name": "Candi Brightwell",
        "case_type": "Committed Intimate Relationship Dissolution",
        "court_name": "Snohomish County Superior Court",
        "facts": (
            "William Miller qualifies as a vulnerable adult under RCW 74.34.020 due to medical vulnerability, PTSD, cognitive impairment, and physical incapacitation. "
            "Evidence shows systematic abuse by Candi Brightwell, including property rights violations, medical emergency exploitation, and business interference. "
            "Court filings and declarations confirm Miller as the primary victim, with Candi as the aggressor using DARVO tactics. "
            "Critical legal issues include lockout during recovery, tool disposal, and withholding of property and business assets. "
            "Legal standards and RCW analysis support modification/termination of protection orders and recognition of Miller's status."
        )
    }

    # Enhanced communications text for NLP analysis
    communications_text = (
        "Text/call transcripts show Candi as primary aggressor. "
        "You need to stop being so dramatic about everything. "
        "Nobody believes your crazy stories anyway. "
        "You're just trying to get attention with all this nonsense. "
        "I'm the only one who puts up with your behavior. "
        "Your family doesn't even want to deal with you anymore. "
        "You better not try to take anything that's mine. "
        "Police reports document Miller's suicidal ideation and garage banishment. "
        "Medical records confirm vulnerable adult status. "
        "Pattern of DARVO manipulation tactics by Candi. "
        "Surveillance admission and EMS records document medical crisis events. "
        "You made me do this by not listening to me. "
        "If you just did what I said, none of this would have happened."
    )

    # 1) Analyze available templates and structure
    print("\n=== TEMPLATE ANALYSIS ===")
    if available_templates:
        template_to_use = None
        for template in available_templates:
            if "protection" in template["name"].lower() or "contempt" in template["name"].lower():
                template_to_use = template
                print(f"[TEMPLATES] Selected: {template['name']}")
                break
        
        if not template_to_use:
            template_to_use = available_templates[0]
            print(f"[TEMPLATES] Using: {template_to_use['name']}")
    else:
        template_to_use = {"path": "templates/motion_template.docx", "name": "generic_template"}
        print("[TEMPLATES] Using generic template")

    # 2) Generate Draft Legal Document with official template
    output_doc = generate_legal_document(
        template_to_use["path"], 
        case_data, 
        f"outputs/{case_data['case_number']}_enhanced_motion.docx"
    )

    # 3) Enhanced Communications Analysis with NLP
    print("\n=== ENHANCED COMMUNICATIONS ANALYSIS ===")
    comms_analysis = analyze_communications_via_ai(communications_text)

    if isinstance(comms_analysis, CommunicationAnalysisResult):
        print("[COMMUNICATIONS ANALYSIS - AI]")
        print(comms_analysis.ai_analysis or "No AI analysis available")

        print("\n[COMMUNICATIONS ANALYSIS - NLP TECHNICAL]")
        nlp_analysis = comms_analysis.nlp_analysis

    elif isinstance(comms_analysis, dict):
        print("[COMMUNICATIONS ANALYSIS - AI]")
        print(comms_analysis.get("ai_analysis", "No AI analysis available"))

        print("\n[COMMUNICATIONS ANALYSIS - NLP TECHNICAL]")
        nlp_analysis = comms_analysis.get("nlp_analysis", {})
    else:
        print("[COMMUNICATIONS ANALYSIS]\n", comms_analysis)
        nlp_analysis = {}

    if nlp_analysis:
        # Display sentiment analysis
        sentiment = nlp_analysis.get("sentiment", {})
        print(f"Sentiment: {sentiment.get('interpretation', 'Unknown')}")

        # Display coercive patterns
        coercive_patterns = nlp_analysis.get("coercive_patterns", [])
        if coercive_patterns:
            print("Coercive Control Patterns Detected:")
            for pattern in coercive_patterns:
                print(f"  - {pattern['category'].title()}: {pattern['severity']} instances")

        # Display psychological indicators
        psych_indicators = nlp_analysis.get("psychological_indicators", [])
        if psych_indicators:
            print("Psychological Manipulation Indicators:")
            for indicator in psych_indicators:
                print(f"  - {indicator['type'].replace('_', ' ').title()}: {indicator['count']} occurrences")

    # 4) Document Structure Analysis
    print("\n=== DOCUMENT STRUCTURE ANALYSIS ===")
    if output_doc:
        structure = analyze_document_structure(output_doc)
        if structure:
            print(f"Document Analysis: {structure['paragraph_count']} paragraphs")
            if structure['heading_structure']:
                print("Document Structure:")
                for heading in structure['heading_structure'][:5]:  # Show first 5 headings
                    print(f"  - {heading['level']}: {heading['text'][:50]}...")

    # 5) Summarize Facts and Generate Arguments (unchanged)
    fact_summary = summarize_case_facts(case_data["facts"])
    arguments = generate_legal_arguments(fact_summary)
    print("\n[CASE SUMMARY]\n", fact_summary)
    print("\n[LEGAL ARGUMENTS]")
    for arg in arguments:
        print("- ", arg)

    # 6) Extract Document Metadata via AI Studio (unchanged)
    extracted_data = ai_studio_extract_metadata([output_doc] if output_doc else [])
    print("\n[AI STUDIO EXTRACTED DATA]\n", json.dumps(extracted_data, indent=2))

    # 7) Expert witness outreach emails (unchanged)
    expert_list = [
        {"type": "psychologist", "provider": "Wilson Psychological & Forensic Services"},
        {"type": "psychologist", "provider": "Snohomish Counseling Collective"},
        {"type": "psychologist", "provider": "Dr. Rachael Silverman"},
        {"type": "financial_forensic", "provider": "4 Corners Financial Forensics"},
        {"type": "financial_forensic", "provider": "Family Law Consulting"},
        {"type": "domestic_violence", "provider": "LCADV Expert Witness Project"}
    ]

    print("\n=== EXPERT WITNESS OUTREACH ===")
    for expert in expert_list:
        email_body = generate_intake_email(expert["type"], expert["provider"])
        send_email(email_body)

    print("\n[SYSTEM] Enhanced Legal Case Workflow Complete!")
    print(f"[SYSTEM] Enhanced document saved: {output_doc}")
    print(f"[SYSTEM] NLP analysis completed with {len(coercive_patterns) if 'coercive_patterns' in locals() else 0} coercive patterns detected")

if __name__ == "__main__":
    main()
